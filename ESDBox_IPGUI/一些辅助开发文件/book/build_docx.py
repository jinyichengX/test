#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
ESDBox_IPGUI 技术书籍 DOCX 生成器

将 book/src/ 目录下的 Markdown 源文件转换为专业排版风格的 DOCX 文档。
使用方法:
    python build_docx.py

输出:
    output/ESDBox_IPGUI_嵌入式图形系统技术手册.docx
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 配置 ────────────────────────────────────────────────
SRC_DIR = os.path.join(os.path.dirname(__file__), "src")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "ESDBox_IPGUI_嵌入式图形系统技术手册.docx")

SOURCE_FILES = [
    "00-preface.md",
    "01-gfx-engine.md",
    "02-composite-system.md",
]

# ── 样式设置 ────────────────────────────────────────────
FONT_CN = "微软雅黑"
FONT_EN = "Consolas"
FONT_BODY_CN = "宋体"
FONT_BODY_EN = "Times New Roman"

def set_font(run, name_cn, name_en, size_pt, bold=False, color=None):
    """设置 run 的中英文字体"""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_styles(doc):
    """添加自定义样式"""
    # 正文样式
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = FONT_BODY_EN
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY_CN)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)

    # 标题样式
    for level, (size, bold, color) in enumerate([
        (22, True, (0x1a, 0x3c, 0x6d)),   # Heading 1
        (17, True, (0x2c, 0x5f, 0x9e)),   # Heading 2
        (14, True, (0x3a, 0x7c, 0xbf)),   # Heading 3
    ], 1):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.size = Pt(size)
        h_style.font.bold = bold
        h_style.font.color.rgb = RGBColor(*color)
        h_style.font.name = FONT_EN
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
        h_style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h_style.paragraph_format.space_after = Pt(8)

def add_horizontal_rule(doc):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

def process_inline_formatting(paragraph, text):
    """处理行内格式：**粗体** `代码` *斜体* [链接]"""
    # 处理粗体 **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # 处理行内代码 `text`
            sub_parts = re.split(r'(`[^`]+`)', part)
            for sp in sub_parts:
                if sp.startswith('`') and sp.endswith('`'):
                    run = paragraph.add_run(sp[1:-1])
                    run.font.name = FONT_EN
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
                else:
                    paragraph.add_run(sp)

def add_code_block(doc, code_lines):
    """添加代码块"""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(line if line else ' ')
        run.font.name = FONT_EN
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers), style='Table Grid')
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    doc.add_paragraph()  # 表后空行
    return table

def convert_markdown_to_docx(doc, md_text):
    """将 Markdown 文本转换为 Word 文档内容"""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_rows = []
    table_headers = []

    while i < len(lines):
        line = lines[i]

        # 代码块处理
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # 表格处理
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
                # 解析表头
                cells = [c.strip() for c in line.split('|')[1:-1]]
                table_headers = cells
                i += 1
                # 跳过分隔行
                if i < len(lines) and '---' in lines[i]:
                    i += 1
                continue
            else:
                cells = [c.strip() for c in line.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
                # 检查下一行是否还是表格
                if i < len(lines) and lines[i].strip().startswith('|'):
                    continue
                else:
                    add_table(doc, table_headers, table_rows)
                    in_table = False
                    continue

        # 结束表格（空行）
        if in_table and not line.strip():
            add_table(doc, table_headers, table_rows)
            in_table = False
            i += 1
            continue

        # 标题
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)

        # 水平线/分隔
        elif line.strip() == '---':
            add_horizontal_rule(doc)

        # 引用块
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(line[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.size = Pt(10)

        # 普通段落（非空）
        elif line.strip():
            # 跳过纯链接引用行
            if line.strip().startswith('[') and line.strip().endswith(')'):
                i += 1
                continue
            p = doc.add_paragraph()
            process_inline_formatting(p, line)

        i += 1

    # 处理文件末尾未关闭的表格/代码块
    if in_code_block and code_buffer:
        add_code_block(doc, code_buffer)
    if in_table and table_rows:
        add_table(doc, table_headers, table_rows)


def build():
    """主构建函数"""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(18.5)
    section.page_height = Cm(26)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)

    add_styles(doc)

    # ── 封面 ──
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ESDBox_IPGUI")
    set_font(run, FONT_CN, FONT_EN, 32, True, (0x1a, 0x3c, 0x6d))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("嵌入式图形系统技术手册")
    set_font(run, FONT_CN, FONT_EN, 26, True, (0x2c, 0x5f, 0x9e))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— 绘图引擎与颜色合成系统深度解析 ——")
    set_font(run, FONT_CN, FONT_EN, 14, False, (0x66, 0x66, 0x66))

    for _ in range(5):
        doc.add_paragraph()

    info_lines = [
        "模块聚焦: core/gfx（图形渲染）与 core/composite（颜色合成）",
        "适用环境: 嵌入式 MCU / 低功耗处理器",
        "语言栈:  纯 C99，无运行时依赖",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_font(run, FONT_CN, FONT_EN, 11, False, (0x88, 0x88, 0x88))

    doc.add_page_break()

    # ── 各章节 ──
    print("[构建] 开始生成 DOCX...")
    for sf in SOURCE_FILES:
        src_path = os.path.join(SRC_DIR, sf)
        if not os.path.exists(src_path):
            print(f"  [警告] 未找到源文件: {src_path}")
            continue

        print(f"  [处理] {sf}")
        with open(src_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        convert_markdown_to_docx(doc, md_content)
        doc.add_page_break()

    # ── 保存 ──
    doc.save(OUTPUT_FILE)
    print(f"\n[完成] 文档已生成: {OUTPUT_FILE}")
    print(f"        文件大小: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")

    return OUTPUT_FILE


if __name__ == "__main__":
    build()
